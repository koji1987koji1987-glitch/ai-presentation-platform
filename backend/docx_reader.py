import sys
import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    # Also extract table cells to capture tabular context
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    return '\n'.join(full_text)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 docx_reader.py <docx_path>")
        sys.exit(1)
    
    path = sys.argv[1]
    try:
        content = read_docx(path)
        print(content)
    except Exception as e:
        print(f"Error reading docx: {e}", file=sys.stderr)
        sys.exit(1)
